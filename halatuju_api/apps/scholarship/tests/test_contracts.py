"""Contract module — authoring immutability, quiz generation (mocked Gemini),
the deployment lifecycle, and the reader seams.

Gemini is mocked exactly like the reports tests (``@patch('google.genai.Client')``
+ overridden ``GEMINI_API_KEY``) — never a live call in CI.
"""
import datetime
import json
from unittest.mock import MagicMock, patch

from django.test import TestCase, override_settings

from apps.courses.models import StudentProfile
from apps.scholarship import contracts
from apps.scholarship.contracts import ContractsError
from apps.scholarship.models import (
    BursaryAgreement, ContractTemplate, ScholarshipApplication, ScholarshipCohort,
)

from apps.scholarship.tests.contract_helpers import (
    brightpath_org, make_deployable, seed_draft,
)

VALID_QUIZ = {'tag': 't', 'plain': 'p', 'question': 'q',
              'options': ['a', 'b', 'c'], 'correct': 1, 'why': 'w'}


def _deploy(template, **kw):
    contracts.submit_for_deployment(template)
    return contracts.deploy(template, is_super=True, **kw)


# ─────────────────────────────────────────────────────────────────────────────
# Authoring
# ─────────────────────────────────────────────────────────────────────────────
class TestCreateTemplate(TestCase):
    def test_creates_a_draft(self):
        t = contracts.create_template(brightpath_org(), '2027-v1', created_by_email='a@b.c')
        self.assertEqual(t.status, 'draft')
        self.assertEqual(t.version, '2027-v1')
        self.assertEqual(t.created_by_email, 'a@b.c')

    def test_blank_version_refused(self):
        with self.assertRaises(ContractsError) as cm:
            contracts.create_template(brightpath_org(), '   ')
        self.assertEqual(cm.exception.code, 'version_required')

    def test_duplicate_version_refused(self):
        contracts.create_template(brightpath_org(), 'dup')
        with self.assertRaises(ContractsError) as cm:
            contracts.create_template(brightpath_org(), 'dup')
        self.assertEqual(cm.exception.code, 'version_exists')

    def test_copy_from_clones_content(self):
        src = seed_draft('2026-v1')
        clone = contracts.create_template(brightpath_org(), '2027-v1', copy_from=src)
        self.assertEqual(clone.clauses.count(), src.clauses.count())
        self.assertEqual(clone.schedule_rows.count(), src.schedule_rows.count())
        self.assertEqual(clone.status, 'draft')
        # Clone starts unvetted/undeployed regardless of the source's stamps.
        self.assertEqual(clone.deployed_by_at, None)


class TestAuthoringIsDraftOnly(TestCase):
    """The immutability guarantee: no authoring call touches a non-draft template."""

    def setUp(self):
        self.t = _deploy(make_deployable())  # now 'active'

    def test_update_config_refused(self):
        with self.assertRaises(ContractsError) as cm:
            contracts.update_config(self.t, counterparty_title='x')
        self.assertEqual(cm.exception.code, 'not_draft')

    def test_replace_clauses_refused(self):
        with self.assertRaises(ContractsError) as cm:
            contracts.replace_clauses(self.t, [])
        self.assertEqual(cm.exception.code, 'not_draft')

    def test_replace_schedule_refused(self):
        with self.assertRaises(ContractsError) as cm:
            contracts.replace_schedule(self.t, [])
        self.assertEqual(cm.exception.code, 'not_draft')

    def test_record_vetting_refused(self):
        with self.assertRaises(ContractsError) as cm:
            contracts.record_vetting(self.t, vetted_by_name='x',
                                     vetted_on=datetime.date(2026, 1, 1),
                                     attested_by_email='a@b.c')
        self.assertEqual(cm.exception.code, 'not_draft')


class TestUpdateConfig(TestCase):
    def test_sets_whitelisted_fields(self):
        t = seed_draft()
        contracts.update_config(t, counterparty_name='ACME', parent_pin_required=False)
        t.refresh_from_db()
        self.assertEqual(t.counterparty_name, 'ACME')
        self.assertFalse(t.parent_pin_required)

    def test_unknown_field_refused(self):
        t = seed_draft()
        with self.assertRaises(ContractsError) as cm:
            contracts.update_config(t, not_a_field='x')
        self.assertEqual(cm.exception.code, 'unknown_config_field')


class TestReplaceSchedule(TestCase):
    def test_duplicate_pathway_variant_refused(self):
        t = seed_draft()
        rows = [
            {'pathway': 'stpm', 'variant': '', 'monthly_amount': '200', 'start_month': 7,
             'paid_offsets': [0]},
            {'pathway': 'stpm', 'variant': '', 'monthly_amount': '200', 'start_month': 7,
             'paid_offsets': [0]},
        ]
        with self.assertRaises(ContractsError) as cm:
            contracts.replace_schedule(t, rows)
        self.assertEqual(cm.exception.code, 'duplicate_schedule_row')


# ─────────────────────────────────────────────────────────────────────────────
# Quiz generation (mocked Gemini)
# ─────────────────────────────────────────────────────────────────────────────
def _mock_genai(text):
    response = MagicMock()
    response.text = text
    client = MagicMock()
    client.models.generate_content.return_value = response
    cls = MagicMock(return_value=client)
    return cls, client


@override_settings(GEMINI_API_KEY='test-key', CONTRACT_QUIZ_MODEL='gemini-2.5-pro')
class TestGenerateQuiz(TestCase):
    def setUp(self):
        self.t = seed_draft()
        self.clause = self.t.clauses.get(order=2)  # a non-candidate to (re)generate

    def _payload(self):
        return json.dumps({'en': VALID_QUIZ, 'ms': VALID_QUIZ, 'ta': VALID_QUIZ})

    @patch('google.genai.Client')
    def test_success_saves_quiz_and_stamps_model(self, mock_cls):
        cls, client = _mock_genai(self._payload())
        mock_cls.side_effect = cls
        contracts.generate_quiz(self.clause)
        self.clause.refresh_from_db()
        self.assertTrue(self.clause.is_quiz_candidate)
        self.assertEqual(self.clause.quiz_en['question'], 'q')
        self.assertEqual(self.clause.quiz_generated_model, 'gemini-2.5-pro')

    @patch('google.genai.Client')
    def test_uses_single_configured_model_no_downgrade(self, mock_cls):
        cls, client = _mock_genai(self._payload())
        mock_cls.side_effect = cls
        contracts.generate_quiz(self.clause)
        # Exactly ONE model call, with the configured model (no cascade).
        self.assertEqual(client.models.generate_content.call_count, 1)
        _, kwargs = client.models.generate_content.call_args
        self.assertEqual(kwargs['model'], 'gemini-2.5-pro')

    @patch('google.genai.Client')
    def test_bad_json_raises(self, mock_cls):
        cls, _ = _mock_genai('not json at all')
        mock_cls.side_effect = cls
        with self.assertRaises(ContractsError) as cm:
            contracts.generate_quiz(self.clause)
        self.assertEqual(cm.exception.code, 'quiz_bad_json')

    @patch('google.genai.Client')
    def test_structurally_invalid_output_raises(self, mock_cls):
        bad = json.dumps({'en': {'options': ['a', 'b'], 'correct': 0}})
        cls, _ = _mock_genai(bad)
        mock_cls.side_effect = cls
        with self.assertRaises(ContractsError) as cm:
            contracts.generate_quiz(self.clause)
        self.assertEqual(cm.exception.code, 'quiz_invalid')

    @patch('google.genai.Client')
    def test_fenced_json_is_tolerated(self, mock_cls):
        cls, _ = _mock_genai('```json\n' + self._payload() + '\n```')
        mock_cls.side_effect = cls
        contracts.generate_quiz(self.clause)
        self.clause.refresh_from_db()
        self.assertEqual(self.clause.quiz_en['correct'], 1)

    def test_missing_api_key_raises(self):
        with override_settings(GEMINI_API_KEY=''):
            with self.assertRaises(ContractsError) as cm:
                contracts.generate_quiz(self.clause)
        self.assertEqual(cm.exception.code, 'quiz_ai_unconfigured')

    def test_generate_on_non_draft_refused(self):
        active = _deploy(make_deployable('2027-x'))
        clause = active.clauses.get(order=2)
        with self.assertRaises(ContractsError) as cm:
            contracts.generate_quiz(clause)
        self.assertEqual(cm.exception.code, 'not_draft')


# ─────────────────────────────────────────────────────────────────────────────
# Lifecycle
# ─────────────────────────────────────────────────────────────────────────────
class TestLifecycle(TestCase):
    def test_submit_moves_draft_to_pending(self):
        t = make_deployable()
        contracts.submit_for_deployment(t, submitted_by_email='a@b.c')
        t.refresh_from_db()
        self.assertEqual(t.status, 'pending_deployment')
        self.assertEqual(t.submitted_by_email, 'a@b.c')

    def test_submit_refuses_when_invalid(self):
        t = seed_draft()  # no counterparty, no attestation → invalid
        with self.assertRaises(ContractsError) as cm:
            contracts.submit_for_deployment(t)
        self.assertEqual(cm.exception.code, 'not_deployable')
        self.assertIn('T1', cm.exception.errors)
        t.refresh_from_db()
        self.assertEqual(t.status, 'draft')

    def test_revert_returns_pending_to_draft(self):
        t = make_deployable()
        contracts.submit_for_deployment(t)
        contracts.revert_to_draft(t)
        t.refresh_from_db()
        self.assertEqual(t.status, 'draft')
        self.assertIsNone(t.submitted_by_at)

    def test_deploy_is_super_only(self):
        t = make_deployable()
        contracts.submit_for_deployment(t)
        with self.assertRaises(ContractsError) as cm:
            contracts.deploy(t, is_super=False)
        self.assertEqual(cm.exception.code, 'deploy_forbidden')

    def test_deploy_requires_pending(self):
        t = make_deployable()  # still draft
        with self.assertRaises(ContractsError) as cm:
            contracts.deploy(t, is_super=True)
        self.assertEqual(cm.exception.code, 'not_pending')

    def test_deploy_activates_and_stamps(self):
        t = _deploy(make_deployable(), deployed_by_email='super@x.c')
        t.refresh_from_db()
        self.assertEqual(t.status, 'active')
        self.assertEqual(t.deployed_by_email, 'super@x.c')
        self.assertIsNotNone(t.deployed_by_at)

    def test_deploy_atomically_archives_previous_active(self):
        first = _deploy(make_deployable('2026-v1'))
        self.assertEqual(first.status, 'active')
        second = _deploy(make_deployable('2026-v2'))
        first.refresh_from_db()
        second.refresh_from_db()
        self.assertEqual(second.status, 'active')
        self.assertEqual(first.status, 'archived')
        self.assertIsNotNone(first.archived_at)
        # Exactly one active template for the org.
        self.assertEqual(
            ContractTemplate.objects.filter(
                organisation=brightpath_org(), status='active').count(), 1)


# ─────────────────────────────────────────────────────────────────────────────
# Readers
# ─────────────────────────────────────────────────────────────────────────────
class TestReaders(TestCase):
    def test_active_template_for(self):
        self.assertIsNone(contracts.active_template_for(brightpath_org()))
        t = _deploy(make_deployable())
        self.assertEqual(contracts.active_template_for(brightpath_org()), t)

    def test_template_for_application_prefers_pinned(self):
        active = _deploy(make_deployable('2026-v1'))
        pinned = ContractTemplate.objects.create(
            organisation=brightpath_org(), version='pinned', status='archived')
        cohort = ScholarshipCohort.objects.create(code='rd', name='B40', year=2026)
        profile = StudentProfile.objects.create(supabase_user_id='rd-1', grades={}, exam_type='spm')
        app = ScholarshipApplication.objects.create(
            cohort=cohort, profile=profile, status='awarded', chosen_pathway='matric')
        BursaryAgreement.objects.create(application=app, version='pinned', template=pinned)
        # The signed agreement's pinned template wins over the org's active one.
        self.assertEqual(contracts.template_for_application(app), pinned)

    def test_quiz_checkpoints_in_order_with_en_fallback(self):
        t = seed_draft()
        cps_en = contracts.quiz_checkpoints(t, 'en')
        self.assertEqual(len(cps_en), 8)
        # ta clause bodies are blank but the quiz has ta content → ta returned.
        cps_ta = contracts.quiz_checkpoints(t, 'ta')
        self.assertEqual(len(cps_ta), 8)
        self.assertTrue(all('question' in cp for cp in cps_ta))

    def test_resolve_locale_clamps_to_available(self):
        t = seed_draft()
        # ms/ta aren't fully translated at template level → clamp to en.
        self.assertEqual(contracts.resolve_locale('ms', t), 'en')
        self.assertEqual(contracts.resolve_locale('en', t), 'en')
        self.assertEqual(contracts.resolve_locale('ta', None), 'en')

    def test_languages_available_is_english_only_for_seed(self):
        t = seed_draft()
        self.assertEqual(t.languages_available, ['en'])

    def test_render_preview_html_has_banner_and_notice(self):
        t = seed_draft()
        html = contracts.render_preview_html(t, 'en')
        self.assertIn('PREVIEW', html)
        self.assertIn('authoritative', html)
        self.assertIn(t.title_en, html)


class TestClauseNumbering(TestCase):
    def test_clause_numbers_three_levels(self):
        self.assertEqual(
            contracts.clause_numbers([0, 1, 1, 2, 2, 0, 1, 2]),
            ['1.', '1.1.', '1.2.', 'I.', 'II.', '2.', '2.1.', 'I.'])

    def test_normalise_forbids_skipping_and_forces_first_zero(self):
        self.assertEqual(contracts.normalise_levels([0, 2, 1, 3, 0, 2]), [0, 1, 1, 2, 0, 1])
        self.assertEqual(contracts.normalise_levels([2, 1]), [0, 1])

    def test_roman(self):
        self.assertEqual([contracts._roman(n) for n in range(1, 6)], ['i', 'ii', 'iii', 'iv', 'v'])


@override_settings(GEMINI_API_KEY='test-key', CONTRACT_QUIZ_MODEL='gemini-2.5-pro')
class TestClauseHierarchy(TestCase):
    def _hierarchy(self, template):
        contracts.replace_clauses(template, [
            {'heading_en': 'Purpose', 'body_en': 'A', 'level': 0, 'is_quiz_candidate': True},
            {'heading_en': 'Eligibility', 'body_en': 'B', 'level': 1, 'is_quiz_candidate': True},
            {'heading_en': 'Instalment', 'body_en': 'C', 'level': 2},
            {'heading_en': 'Obligations', 'body_en': 'D', 'level': 0},
        ])
        return list(template.clauses.order_by('order'))

    def test_replace_clauses_stores_level_and_restricts_quiz_to_top_level(self):
        clauses = self._hierarchy(seed_draft())
        self.assertEqual([c.level for c in clauses], [0, 1, 2, 0])
        # the level-1 clause's quiz flag was dropped (only level-0 carries a quiz)
        self.assertTrue(clauses[0].is_quiz_candidate)
        self.assertFalse(clauses[1].is_quiz_candidate)

    def test_replace_clauses_normalises_a_skip(self):
        d = seed_draft()
        contracts.replace_clauses(d, [
            {'heading_en': 'A', 'body_en': 'a', 'level': 0},
            {'heading_en': 'B', 'body_en': 'b', 'level': 2},  # skip → normalised to 1
        ])
        self.assertEqual([c.level for c in d.clauses.order_by('order')], [0, 1])

    def test_clause_and_descendants_covers_subtree(self):
        clauses = self._hierarchy(seed_draft())
        top, sub, subsub, other = clauses
        self.assertEqual([c.pk for c in contracts._clause_and_descendants(top)],
                         [top.pk, sub.pk, subsub.pk])
        self.assertEqual([c.pk for c in contracts._clause_and_descendants(other)], [other.pk])

    @patch('google.genai.Client')
    def test_generate_quiz_refuses_non_top_level(self, mock_cls):
        cls, _ = _mock_genai(json.dumps({'en': VALID_QUIZ, 'ms': VALID_QUIZ, 'ta': VALID_QUIZ}))
        mock_cls.side_effect = cls
        sub = self._hierarchy(seed_draft())[1]   # level 1
        with self.assertRaises(ContractsError) as cm:
            contracts.generate_quiz(sub)
        self.assertEqual(cm.exception.code, 'quiz_not_top_level')

    @patch('google.genai.Client')
    def test_quiz_prompt_includes_subtree_text(self, mock_cls):
        cls, client = _mock_genai(json.dumps({'en': VALID_QUIZ, 'ms': VALID_QUIZ, 'ta': VALID_QUIZ}))
        mock_cls.side_effect = cls
        top = self._hierarchy(seed_draft())[0]
        contracts.generate_quiz(top)
        _, kwargs = client.models.generate_content.call_args
        prompt = kwargs['contents']
        # the top clause's quiz prompt carries its sub-clause bodies too
        self.assertIn('Eligibility', prompt)
        self.assertIn('Instalment', prompt)

    @patch('apps.scholarship.contracts._docx_structure', return_value=None)
    @patch('apps.scholarship.contracts._extract_docx_text', return_value='doc text')
    @patch('google.genai.Client')
    def test_segment_docx_gemini_fallback_returns_normalised_levels(
            self, mock_cls, _extract, _struct):
        # An unstyled doc (_docx_structure → None) falls back to Gemini; the proposed
        # levels are still normalised (no skipping) and returned under 'clauses'.
        segments = json.dumps([
            {'heading': 'A', 'body': 'a', 'level': 0},
            {'heading': 'B', 'body': 'b', 'level': 2},   # skip → 1
            {'heading': 'C', 'body': 'c', 'level': 1},
        ])
        cls, _ = _mock_genai(segments)
        mock_cls.side_effect = cls
        got = contracts.segment_docx(b'x')
        self.assertEqual([c['level'] for c in got['clauses']], [0, 1, 1])
        self.assertEqual(got['title'], '')
        self.assertEqual(got['preamble'], '')

    def test_render_shows_computed_numbers(self):
        from apps.scholarship import bursary
        d = seed_draft('2027-hier-d')
        self._hierarchy(d)
        cohort = ScholarshipCohort.objects.create(
            code='hc', name='B40', year=2026, owning_organisation=d.organisation)
        prof = StudentProfile.objects.create(supabase_user_id='hs', name='S', nric='000101-10-1233')
        app = ScholarshipApplication.objects.create(
            cohort=cohort, profile=prof, status='awarded', notify_email='s@e.test', award_amount=3000)
        p = {'award_amount': 3000, 'payment_schedule': 'x', 'institution_name': 'U',
             'course_name': 'N', 'commencement_date': None, 'progress_standard': 'Pass',
             'foundation_signatory_name': 'F', 'foundation_signatory_title': 't',
             'foundation_signatory_nric': ''}
        html = bursary.render_agreement_html(
            app, p, student={'name': 'S', 'nric': 'x', 'signed_at': None},
            guarantor={'name': 'G', 'nric': 'y', 'relationship': 'father', 'signed_at': None},
            foundation={'name': 'F', 'title': 't', 'nric': '', 'signed_at': None},
            witness={'by': '', 'org': '', 'signed_at': None}, template=d)
        self.assertIn('>1.<', html)      # top-level number (bold)
        self.assertIn('>1.1.<', html)    # sub-clause (Word style, trailing stop)
        self.assertIn('>I.<', html)      # sub-sub-clause (uppercase roman)


# ─────────────────────────────────────────────────────────────────────────────
# Word import — deterministic structural parse (2026-07-21)
# ─────────────────────────────────────────────────────────────────────────────
def _list_para(para, num_id, ilvl):
    """Attach Word list numbering (numId/ilvl) to a paragraph — the numbering that
    _docx_structure reads and that paragraph.text does NOT contain."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = para._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl_el = OxmlElement('w:ilvl'); ilvl_el.set(qn('w:val'), str(ilvl))
    numId_el = OxmlElement('w:numId'); numId_el.set(qn('w:val'), str(num_id))
    numPr.append(ilvl_el); numPr.append(numId_el)
    pPr.append(numPr)
    return para


def _sample_agreement_bytes():
    """A .docx mirroring the real 'Donor–Student Conditional Agreement' structure:
    Title, a preamble with hand-typed brackets, two top clauses each with a full-sentence
    sub-clause and a couple of roman list items."""
    import io as _io
    from docx import Document
    doc = Document()
    doc.add_paragraph('Donor–Student Conditional Agreement', style='Title')
    doc.add_paragraph(
        'This Agreement is made between [Student Full Name & NRIC] of [Student Address].')
    _list_para(doc.add_paragraph('Definitions and Interpretation', style='Heading 1'), 6, 0)
    _list_para(doc.add_paragraph(
        'In this Agreement the following expressions have the meanings below.',
        style='Heading 1'), 6, 1)
    _list_para(doc.add_paragraph('"Agreement" means this agreement.'), 4, 0)
    _list_para(doc.add_paragraph('"Bursary" means the sum awarded.'), 4, 0)
    _list_para(doc.add_paragraph('Student Obligations', style='Heading 1'), 6, 0)
    _list_para(doc.add_paragraph('The Student agrees to:', style='Heading 1'), 6, 1)
    _list_para(doc.add_paragraph('attend all classes.'), 8, 0)
    _list_para(doc.add_paragraph('maintain the progress standard.'), 8, 0)
    buf = _io.BytesIO(); doc.save(buf)
    return buf.getvalue()


class TestDocxStructure(TestCase):
    def test_parses_hierarchy_from_word_numbering(self):
        got = contracts.segment_docx(_sample_agreement_bytes())
        self.assertEqual([c['level'] for c in got['clauses']],
                         [0, 1, 2, 2, 0, 1, 2, 2])
        headings = [c['heading'] for c in got['clauses']]
        self.assertEqual(headings[0], 'Definitions and Interpretation')
        self.assertEqual(headings[4], 'Student Obligations')
        # roman-level items carry the text in the body, not the heading
        self.assertEqual(got['clauses'][2]['heading'], '')
        self.assertIn('"Agreement" means', got['clauses'][2]['body'])

    def test_captures_title_and_preamble(self):
        got = contracts.segment_docx(_sample_agreement_bytes())
        self.assertEqual(got['title'], 'Donor–Student Conditional Agreement')
        self.assertIn('This Agreement is made', got['preamble'])

    def test_converts_bracket_placeholders_to_tokens(self):
        got = contracts.segment_docx(_sample_agreement_bytes())
        # a recognised bracket → token; an unrecognised one is left verbatim
        self.assertIn('{{student_name}} ({{student_nric}})', got['preamble'])
        self.assertIn('[Student Address]', got['preamble'])

    def test_unstyled_doc_yields_none_from_structure(self):
        import io as _io
        from docx import Document
        doc = Document()
        doc.add_paragraph('1. Some clause with a hand-typed number and no styles.')
        doc.add_paragraph('2. Another one.')
        buf = _io.BytesIO(); doc.save(buf)
        self.assertIsNone(contracts._docx_structure(buf.getvalue()))


class TestSubstituteVars(TestCase):
    def test_fills_known_leaves_unknown(self):
        ctx = {'student_name': 'Aisha', 'institution': ''}
        out = contracts.substitute_vars(
            'For {{student_name}} at {{institution}} — {{mystery}}.', ctx)
        # known-with-value filled; known-but-empty → blank; unknown → verbatim
        self.assertEqual(out, 'For Aisha at  — {{mystery}}.')

    def test_tolerates_whitespace_and_empty(self):
        self.assertEqual(contracts.substitute_vars('{{ student_name }}',
                                                   {'student_name': 'B'}), 'B')
        self.assertEqual(contracts.substitute_vars('', {'x': 'y'}), '')


class TestContractRenderRich(TestCase):
    """Bold (**…**) and {{variable}} substitution in the rendered agreement."""
    def _render(self, body_en):
        from apps.scholarship import bursary
        d = seed_draft('2027-rich-d')
        contracts.replace_clauses(d, [{'heading_en': 'Terms', 'body_en': body_en, 'level': 0}])
        cohort = ScholarshipCohort.objects.create(
            code='rc', name='B40', year=2026, owning_organisation=d.organisation)
        prof = StudentProfile.objects.create(
            supabase_user_id='rs', name='S', nric='000101-10-1233')
        app = ScholarshipApplication.objects.create(
            cohort=cohort, profile=prof, status='awarded',
            notify_email='s@e.test', award_amount=3000)
        p = {'award_amount': 3000, 'payment_schedule': 'x', 'institution_name': 'Universiti X',
             'course_name': 'N', 'commencement_date': None, 'progress_standard': 'Pass',
             'foundation_signatory_name': 'F', 'foundation_signatory_title': 't',
             'foundation_signatory_nric': ''}
        return bursary.render_agreement_html(
            app, p, student={'name': 'Aisha Binti Ali', 'nric': 'x', 'signed_at': None},
            guarantor={'name': 'G', 'nric': 'y', 'relationship': 'father', 'signed_at': None},
            foundation={'name': 'F', 'title': 't', 'nric': '', 'signed_at': None},
            witness={'by': '', 'org': '', 'signed_at': None}, template=d)

    def test_bold_marker_becomes_b_tag(self):
        html = self._render('The **Bursary** is conditional.')
        self.assertIn('<b>Bursary</b>', html)
        self.assertNotIn('**Bursary**', html)

    def test_variable_is_substituted(self):
        html = self._render('Awarded to {{student_name}} at {{institution}}.')
        self.assertIn('Aisha Binti Ali', html)
        self.assertIn('Universiti X', html)
        self.assertNotIn('{{student_name}}', html)

    def test_unknown_variable_left_visible(self):
        html = self._render('Ref {{not_a_real_var}} here.')
        self.assertIn('{{not_a_real_var}}', html)


class TestPreviewRender(TestCase):
    """render_preview_html — escaping, hierarchical numbering, and **bold** (TD-163: the
    preview must match the signed agreement, and must not corrupt on HTML-special chars)."""
    def _draft(self):
        d = seed_draft('2027-prev-r')
        contracts.replace_clauses(d, [
            {'heading_en': 'Terms & Conditions', 'body_en': 'The **Bursary** is a gift.', 'level': 0},
            {'heading_en': 'Sub-point', 'body_en': 'Amount < 3000 & > 0.', 'level': 1},
        ])
        return d

    def test_escapes_html_special_chars(self):
        html = contracts.render_preview_html(self._draft(), 'en')
        self.assertIn('Terms &amp; Conditions', html)   # & escaped
        self.assertIn('&lt; 3000 &amp; &gt; 0', html)   # < > & escaped in body
        self.assertNotIn('Terms & Conditions', html)    # never the raw ampersand

    def test_hierarchical_numbering(self):
        html = contracts.render_preview_html(self._draft(), 'en')
        self.assertIn('>1.<', html)      # top-level
        self.assertIn('>1.1.<', html)    # sub-clause (Word style, NOT the old flat order '2.')

    def test_bold_marker_rendered(self):
        html = contracts.render_preview_html(self._draft(), 'en')
        self.assertIn('<b>Bursary</b>', html)
        self.assertNotIn('**Bursary**', html)


class TestHeadingOverflow(TestCase):
    """Guards for the varchar(255) heading column — a full sub-clause the author styled as a
    Heading must not 500 the clause save (regression: import 500 / empty upload, 2026-07-21)."""
    def test_long_heading_paragraph_parses_as_body(self):
        import io as _io
        from docx import Document
        doc = Document()
        _list_para(doc.add_paragraph('Definitions', style='Heading 1'), 6, 0)
        long_sentence = 'The Donor agrees to provide the Student with a bursary ' + ('x ' * 200) + '.'
        _list_para(doc.add_paragraph(long_sentence, style='Heading 1'), 6, 1)
        buf = _io.BytesIO(); doc.save(buf)
        got = contracts.segment_docx(buf.getvalue())
        self.assertEqual(got['clauses'][0]['heading'], 'Definitions')   # short stays a title
        self.assertEqual(got['clauses'][1]['heading'], '')              # long → body
        self.assertIn('The Donor agrees', got['clauses'][1]['body'])
        self.assertTrue(all(len(c['heading']) <= 255 for c in got['clauses']))

    def test_replace_clauses_folds_overlong_heading_into_body(self):
        # The hard save guard: even a hand-typed 300-char heading must not overflow the column.
        d = seed_draft('2027-overflow')
        longh = 'A very long clause heading that should live in the body instead ' + ('y ' * 150)
        self.assertGreater(len(longh), 255)
        contracts.replace_clauses(d, [{'heading_en': longh, 'body_en': 'tail.', 'level': 0}])
        c = d.clauses.first()
        self.assertEqual(c.heading_en, '')          # moved out of the 255 column
        self.assertIn('A very long clause heading', c.body_en)
        self.assertIn('tail.', c.body_en)           # original body preserved

    def test_short_heading_unaffected(self):
        d = seed_draft('2027-shorth')
        contracts.replace_clauses(d, [{'heading_en': 'Definitions', 'body_en': 'x', 'level': 0}])
        c = d.clauses.first()
        self.assertEqual(c.heading_en, 'Definitions')
        self.assertEqual(c.body_en, 'x')


class TestSharedClauseRender(TestCase):
    """render_clauses_html — hanging-indent table, bold ONLY at level 0, Word numbering."""
    def test_bold_rules_and_word_numbering(self):
        html = contracts.render_clauses_html([
            ('Definitions', '', 0),
            ('In this Agreement', '', 1),
            ('', 'Agreement means this document.', 2),
        ])
        # level 0: bold number AND bold heading
        self.assertIn('<b>1.</b>', html)
        self.assertIn('<b>Definitions</b>', html)
        # level 1: Word style number (NOT bold); heading not bold
        self.assertIn('>1.1.</td>', html)
        self.assertNotIn('<b>1.1.</b>', html)
        self.assertNotIn('<b>In this Agreement</b>', html)
        # level 2: uppercase roman + full stop, not bold; body in the text cell (hanging indent)
        self.assertIn('>I.</td>', html)
        self.assertNotIn('<b>I.</b>', html)
        self.assertIn('Agreement means this document.', html)
        # a per-clause table drives the hanging indent
        self.assertIn('<table', html)

    def test_escapes_and_honours_bold_marker(self):
        html = contracts.render_clauses_html([('R&D terms', 'a **strong** point', 0)])
        self.assertIn('R&amp;D terms', html)     # escaped
        self.assertIn('<b>strong</b>', html)     # **..** honoured


class TestCounterpartyExtraction(TestCase):
    def test_extracts_name_nric_address(self):
        preamble = ('This Agreement is made on [Date] between Jane Doe, NRIC 800101-01-1234, '
                    'of 12 Jalan Test, 50000 KL ("Donor"), and {{student_name}}.')
        cp = contracts._extract_counterparty(preamble)
        self.assertEqual(cp['name'], 'Jane Doe')
        self.assertEqual(cp['nric'], '800101-01-1234')
        self.assertIn('12 Jalan Test', cp['address'])
        self.assertNotIn('Donor', cp['address'])   # address stops before the ("Donor") marker

    def test_no_match_returns_blanks(self):
        self.assertEqual(contracts._extract_counterparty('No recital here.'),
                         {'name': '', 'nric': '', 'address': ''})

    def test_config_accepts_counterparty_address(self):
        d = seed_draft('2027-cp-addr')
        contracts.update_config(d, counterparty_address='12 Jalan Test, 50000 KL')
        d.refresh_from_db()
        self.assertEqual(d.counterparty_address, '12 Jalan Test, 50000 KL')


class TestDonorNameToVariable(TestCase):
    """Import replaces the donor's literal name with {{donor_name}} (defined once in Config)."""
    def _bytes(self):
        import io as _io
        from docx import Document
        doc = Document()
        doc.add_paragraph('Bursary Agreement', style='Title')
        doc.add_paragraph(
            'This Agreement is made between John Smith, NRIC 900101-01-1234, of 5 Main St '
            '("Donor"), and {{student_name}}.')
        _list_para(doc.add_paragraph('Definitions', style='Heading 1'), 6, 0)
        _list_para(doc.add_paragraph('"Donor" means John Smith.'), 4, 0)
        buf = _io.BytesIO(); doc.save(buf)
        return buf.getvalue()

    def test_literal_donor_name_becomes_token(self):
        got = contracts.segment_docx(self._bytes())
        self.assertEqual(got['counterparty']['name'], 'John Smith')       # kept for the Config field
        self.assertNotIn('John Smith', got['preamble'])                    # swapped out of the text
        self.assertIn('{{donor_name}}', got['preamble'])
        donor_clause = next(c for c in got['clauses'] if 'means' in c['body'])
        self.assertEqual(donor_clause['body'], '"Donor" means {{donor_name}}.')

    def test_literal_donor_nric_and_address_become_tokens(self):
        got = contracts.segment_docx(self._bytes())
        self.assertEqual(got['counterparty']['nric'], '900101-01-1234')   # kept for Config
        self.assertEqual(got['counterparty']['address'], '5 Main St')
        # both the NRIC and the address are tokenised in the preamble text
        self.assertIn('{{donor_nric}}', got['preamble'])
        self.assertIn('{{donor_address}}', got['preamble'])
        self.assertNotIn('900101-01-1234', got['preamble'])
        self.assertNotIn('5 Main St', got['preamble'])
