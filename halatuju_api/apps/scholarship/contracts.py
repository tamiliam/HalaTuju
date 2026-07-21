"""Contract module service — org-owned, versioned bursary-agreement templates.

Mirrors ``payments.py`` in shape (a ``ContractsError`` carrying a machine code;
module-level service functions over the models). The lifecycle is DEPLOYMENT,
not approval::

    draft → pending_deployment → active → archived

Authoring calls all refuse a non-draft template (``not_draft``) — the
immutability guarantee that lets a signed ``BursaryAgreement`` PROTECT-reference
the exact version it was rendered from, forever. Deploying a new version
atomically archives the org's previous active version.

Sprint 1 scope: model + service + seed + tests. The module is INERT — neither
``bursary.py`` nor ``payments.py`` reads it yet (that cutover is Sprint 2). No
live Gemini / Drive / storage calls happen here; the Gemini seam
(``_gemini_generate``) is mocked in tests exactly like the reports tests.
"""

import json
import re
from decimal import Decimal

from django.conf import settings
from django.db import transaction
from django.utils import timezone

from . import award
from .models import ContractClause, ContractTemplate, PaymentScheduleRow

# The three languages the module carries. English is authoritative.
LANGUAGES = ('en', 'ms', 'ta')

# Config fields an author may set on a draft (whitelist for update_config).
_CONFIG_FIELDS = (
    'title_en', 'title_ms', 'title_ta',
    'preamble_en', 'preamble_ms', 'preamble_ta',
    'progress_standard_en', 'progress_standard_ms', 'progress_standard_ta',
    'counterparty_name', 'counterparty_title', 'counterparty_nric',
    'counterparty_notify_emails', 'parent_role', 'parent_pin_required',
    'witness_policy',
)

# Clause hierarchy (2026-07-19): three levels — 0 clause, 1 sub-clause, 2 sub-sub-clause.
MAX_CLAUSE_LEVEL = 2
_ROMAN = ((10, 'x'), (9, 'ix'), (5, 'v'), (4, 'iv'), (1, 'i'))

# ContractClause.heading_* is CharField(max_length=255); body_* is TextField (unlimited).
# HEADING_MAX_LEN is the hard DB limit (the save guard); a "heading" longer than
# HEADING_TITLE_MAX is treated as a SENTENCE (clause body), not a short title — real docs
# style full sub-clauses as a Heading, which would otherwise overflow the column (TD, 2026-07-21).
HEADING_MAX_LEN = 255
HEADING_TITLE_MAX = 120


def _fit_heading(heading, body):
    """Keep a clause heading inside the varchar(255) column: a heading longer than the limit
    is not a title — fold it into the body (lossless) so a save can NEVER overflow, whatever
    the source (import, Gemini, hand-typed, copy-from). Returns (heading, body)."""
    heading = (heading or '').strip()
    body = body or ''
    if len(heading) <= HEADING_MAX_LEN:
        return heading, body
    return '', (heading + '\n\n' + body).strip() if body else heading


def _roman(n):
    """Lowercase roman numeral for n >= 1 (i, ii, iii, iv, v, …). Small n only (clause depth)."""
    out = []
    for value, sym in _ROMAN:
        while n >= value:
            out.append(sym)
            n -= value
    return ''.join(out)


def normalise_levels(levels):
    """Clamp each level to 0..MAX and forbid SKIPPING a level: a clause may be at most one level
    deeper than the clause before it (the first is forced to 0); going shallower is unrestricted.
    Returns the cleaned list — the single place the no-skip rule lives (both import and save use it)."""
    out = []
    prev = -1  # first allowed level is 0
    for lv in levels:
        try:
            lv = int(lv)
        except (TypeError, ValueError):
            lv = 0
        lv = max(0, min(MAX_CLAUSE_LEVEL, lv))
        if lv > prev + 1:
            lv = prev + 1
        out.append(lv)
        prev = lv
    return out


def clause_numbers(levels):
    """Display label per clause from its level run: level 0 → '1.', '2.'; level 1 → '1.1', '1.2';
    level 2 → 'i)', 'ii)'. Levels are assumed already valid (see normalise_levels). This is the
    ONE numbering source of truth — the FE mirror in lib/clauseNumbering.ts must match it (paired
    test). Numbers are computed, never stored, so reorder/insert always renumbers correctly."""
    counters = [0, 0, 0]
    labels = []
    for lv in levels:
        lv = max(0, min(MAX_CLAUSE_LEVEL, int(lv)))
        counters[lv] += 1
        for deeper in range(lv + 1, MAX_CLAUSE_LEVEL + 1):
            counters[deeper] = 0
        if lv == 0:
            labels.append(f'{counters[0]}.')
        elif lv == 1:
            labels.append(f'{counters[0]}.{counters[1]}')
        else:
            labels.append(f'{_roman(counters[2])})')
    return labels


# ─────────────────────────────────────────────────────────────────────────────
# Merge variables (2026-07-21): {{token}} placeholders an author writes into a clause
# / heading / preamble, resolved to the case's real values at RENDER time (never stored
# resolved — the template keeps the generic token, the signed snapshot gets the value).
# The registry is the single source of truth the editor's "Insert variable" menu reads.
# ─────────────────────────────────────────────────────────────────────────────
CONTRACT_VARS = (
    ('student_name', "The student's full name"),
    ('student_nric', "The student's NRIC"),
    ('guarantor_name', "The guarantor / parent-guardian's name"),
    ('guarantor_relationship', 'The guarantor relationship (e.g. father)'),
    ('donor_name', 'The donor / sponsor name (from the template config)'),
    ('amount', 'The bursary amount (e.g. RM3,000)'),
    ('institution', 'The institution name'),
    ('course', 'The course / programme name'),
    ('commencement_date', 'The course commencement date'),
    ('progress_standard', 'The academic progress standard'),
)
_VAR_RE = re.compile(r'\{\{\s*(\w+)\s*\}\}')


def substitute_vars(text, context):
    """Replace ``{{token}}`` placeholders with values from ``context``. A KNOWN token with
    no value renders empty; an UNKNOWN token is left VERBATIM — a stray ``{{typo}}`` stays
    visible in the document (obviously fixable) rather than becoming a silent blank."""
    if not text:
        return text

    def _repl(m):
        key = m.group(1)
        if key in context:
            val = context[key]
            return '' if val is None else str(val)
        return m.group(0)

    return _VAR_RE.sub(_repl, text)


# Bracket placeholders authors hand-type in Word (e.g. "[Student Full Name & NRIC]") →
# our tokens, applied on import. Only mappings with reliable backing data are listed;
# an unrecognised bracket (e.g. "[Date]", "[Student Address]") is left untouched for the
# author to resolve, never converted into a dangling token.
_BRACKET_VARS = {
    'student full name & nric': '{{student_name}} ({{student_nric}})',
    'student full name and nric': '{{student_name}} ({{student_nric}})',
    'student name & nric': '{{student_name}} ({{student_nric}})',
    'student full name': '{{student_name}}',
    'student name': '{{student_name}}',
    'student nric': '{{student_nric}}',
    'institution': '{{institution}}',
    'course': '{{course}}',
}
_BRACKET_RE = re.compile(r'\[([^\[\]]+)\]')


def _brackets_to_vars(text):
    """Convert an author's hand-typed ``[Bracket Placeholder]`` to the matching ``{{token}}``
    (import-time sugar). Unrecognised brackets are preserved verbatim."""
    if not text:
        return text
    return _BRACKET_RE.sub(
        lambda m: _BRACKET_VARS.get(m.group(1).strip().lower(), m.group(0)), text)


class ContractsError(Exception):
    """Raised by the service with a machine code for the view (e.g. 'not_draft',
    'version_exists', 'deploy_forbidden', 'not_deployable')."""

    def __init__(self, code, message=''):
        self.code = code
        super().__init__(message or code)


class ValidationResult:
    """Outcome of ``validate_for_deployment``: deduped, sorted error/warning
    code lists. ``ok`` is True when there are no errors."""

    def __init__(self, errors, warnings):
        self.errors = sorted(set(errors))
        self.warnings = sorted(set(warnings))

    @property
    def ok(self):
        return not self.errors

    def __repr__(self):
        return f'ValidationResult(errors={self.errors}, warnings={self.warnings})'


def _require_draft(template):
    if template.status != 'draft':
        raise ContractsError('not_draft')


# ─────────────────────────────────────────────────────────────────────────────
# Authoring (draft-only)
# ─────────────────────────────────────────────────────────────────────────────
@transaction.atomic
def create_template(organisation, version, *, created_by_email='', copy_from=None):
    """Create a new DRAFT template for ``organisation``. With ``copy_from`` (a
    ContractTemplate), clone its config + clauses + schedule — but never its
    lifecycle/attestation stamps (a new version starts unvetted, undeployed)."""
    version = (version or '').strip()
    if not version:
        raise ContractsError('version_required')
    if ContractTemplate.objects.filter(organisation=organisation, version=version).exists():
        raise ContractsError('version_exists')
    template = ContractTemplate.objects.create(
        organisation=organisation, version=version, status='draft',
        created_by_email=created_by_email or '',
    )
    if copy_from is not None:
        _clone_content(copy_from, template)
    return template


def _clone_content(source, target):
    for field in _CONFIG_FIELDS:
        setattr(target, field, getattr(source, field))
    target.save()
    clauses = [
        {
            'level': c.level,
            'heading_en': c.heading_en, 'heading_ms': c.heading_ms, 'heading_ta': c.heading_ta,
            'body_en': c.body_en, 'body_ms': c.body_ms, 'body_ta': c.body_ta,
            'is_quiz_candidate': c.is_quiz_candidate,
            'quiz_en': c.quiz_en, 'quiz_ms': c.quiz_ms, 'quiz_ta': c.quiz_ta,
            'quiz_generated_model': c.quiz_generated_model,
        }
        for c in source.clauses.all().order_by('order')
    ]
    replace_clauses(target, clauses)
    rows = [
        {
            'pathway': r.pathway, 'variant': r.variant,
            'label_en': r.label_en, 'label_ms': r.label_ms, 'label_ta': r.label_ta,
            'monthly_amount': r.monthly_amount, 'start_month': r.start_month,
            'paid_offsets': r.paid_offsets, 'sort_order': r.sort_order,
        }
        for r in source.schedule_rows.all()
    ]
    replace_schedule(target, rows)


def update_config(template, **fields):
    """Set whitelisted config fields on a draft. Unknown keys raise."""
    _require_draft(template)
    changed = []
    for key, value in fields.items():
        if key not in _CONFIG_FIELDS:
            raise ContractsError('unknown_config_field', key)
        setattr(template, key, value)
        changed.append(key)
    if changed:
        template.save(update_fields=changed + ['updated_at'])
    return template


@transaction.atomic
def replace_clauses(template, clauses):
    """Atomic PUT of the clause list. Order is assigned by position (contiguous
    1..N by construction). Each item is a flat dict of model fields; a supplied
    quiz payload must be a dict."""
    _require_draft(template)
    template.clauses.all().delete()
    # Levels are normalised across the whole list (clamp 0..MAX, no skipping) — the tree is the
    # (order, level) run. A quiz may only live on a level-0 clause; a sub-clause carrying a quiz
    # flag/payload is dropped (its content is folded into its parent clause's quiz subtree).
    levels = normalise_levels([item.get('level', 0) for item in clauses])
    created = []
    for index, (item, level) in enumerate(zip(clauses, levels), start=1):
        for lang in LANGUAGES:
            payload = item.get(f'quiz_{lang}') or {}
            if payload and not isinstance(payload, dict):
                raise ContractsError('bad_quiz_payload', f'clause {index} quiz_{lang}')
        is_l0 = level == 0
        # An over-long heading (any language) folds into its body so the save can never
        # overflow heading_*'s varchar(255) — the guard for every write path (2026-07-21).
        h_en, b_en = _fit_heading(item.get('heading_en', ''), item.get('body_en', ''))
        h_ms, b_ms = _fit_heading(item.get('heading_ms', ''), item.get('body_ms', ''))
        h_ta, b_ta = _fit_heading(item.get('heading_ta', ''), item.get('body_ta', ''))
        created.append(ContractClause(
            template=template, order=index, level=level,
            heading_en=h_en, heading_ms=h_ms, heading_ta=h_ta,
            body_en=b_en, body_ms=b_ms, body_ta=b_ta,
            is_quiz_candidate=bool(item.get('is_quiz_candidate')) and is_l0,
            quiz_en=(item.get('quiz_en') or {}) if is_l0 else {},
            quiz_ms=(item.get('quiz_ms') or {}) if is_l0 else {},
            quiz_ta=(item.get('quiz_ta') or {}) if is_l0 else {},
            quiz_generated_model=(item.get('quiz_generated_model', '') or '') if is_l0 else '',
        ))
    ContractClause.objects.bulk_create(created)
    return template.clauses.all()


@transaction.atomic
def replace_schedule(template, rows):
    """Atomic PUT of the schedule rows. Rejects a duplicate (pathway, variant)
    in the input before the DB constraint would."""
    _require_draft(template)
    seen = set()
    template.schedule_rows.all().delete()
    created = []
    for index, item in enumerate(rows):
        pathway = (item.get('pathway', '') or '').strip().lower()
        variant = (item.get('variant', '') or '').strip().lower()
        if not pathway:
            raise ContractsError('bad_schedule_row', f'row {index}: missing pathway')
        key = (pathway, variant)
        if key in seen:
            raise ContractsError('duplicate_schedule_row', f'{pathway}/{variant}')
        seen.add(key)
        created.append(PaymentScheduleRow(
            template=template, pathway=pathway, variant=variant,
            label_en=item.get('label_en', '') or '',
            label_ms=item.get('label_ms', '') or '',
            label_ta=item.get('label_ta', '') or '',
            monthly_amount=Decimal(str(item.get('monthly_amount', '0'))),
            start_month=int(item.get('start_month', 0)),
            paid_offsets=list(item.get('paid_offsets') or []),
            sort_order=int(item.get('sort_order', index)),
        ))
    PaymentScheduleRow.objects.bulk_create(created)
    return template.schedule_rows.all()


def record_vetting(template, *, vetted_by_name, vetted_on, attested_by_email):
    """Record the lawyer-vetting attestation (who + date) — the T2 gate."""
    _require_draft(template)
    if not (vetted_by_name and vetted_on and attested_by_email):
        raise ContractsError('incomplete_vetting')
    template.vetted_by_name = vetted_by_name
    template.vetted_on = vetted_on
    template.vetting_attested_by_email = attested_by_email
    template.vetting_attested_at = timezone.now()
    template.save(update_fields=[
        'vetted_by_name', 'vetted_on', 'vetting_attested_by_email',
        'vetting_attested_at', 'updated_at',
    ])
    return template


# ─────────────────────────────────────────────────────────────────────────────
# Quiz generation (Gemini — draft-only, on-demand, single model, no downgrade)
# ─────────────────────────────────────────────────────────────────────────────
def _gemini_generate(prompt, model):
    """The single-model Gemini call — the mockable seam (patched in tests, never
    a live call in CI). Owner decision 4: NO downgrade fallback — if the
    configured model is unconfigured/unavailable/errors, raise; do not degrade."""
    api_key = getattr(settings, 'GEMINI_API_KEY', '')
    if not api_key:
        raise ContractsError('quiz_ai_unconfigured')
    try:
        from google import genai
    except ImportError:
        raise ContractsError('quiz_ai_unavailable')
    client = genai.Client(api_key=api_key)
    response = client.models.generate_content(model=model, contents=prompt)
    return response.text


def _clause_and_descendants(clause):
    """A clause plus its descendant sub-clauses: the following clauses whose level is DEEPER,
    up to (but not including) the next clause at the same-or-shallower level. Lets a top-level
    clause's comprehension quiz cover its entire subtree (sub- and sub-sub-clauses included)."""
    ordered = list(clause.template.clauses.order_by('order'))
    result = []
    collecting = False
    for c in ordered:
        if c.pk == clause.pk:
            collecting = True
            result.append(c)
        elif collecting:
            if c.level > clause.level:
                result.append(c)
            else:
                break
    return result or [clause]


def _build_quiz_prompt(clause):
    """Strict-JSON prompt asking for a comprehension checkpoint in en/ms/ta, built from the whole
    clause SUBTREE (the clause + its sub-clauses) so the checkpoint covers the clause's key points,
    using each language's translations when present, else English."""
    subtree = _clause_and_descendants(clause)

    def block(lang):
        parts = []
        for c in subtree:
            heading = getattr(c, f'heading_{lang}') or c.heading_en
            body = getattr(c, f'body_{lang}') or c.body_en
            seg = f'{heading}\n{body}'.strip()
            if seg:
                parts.append(seg)
        return f'[{lang}]\n' + '\n\n'.join(parts)

    clause_text = '\n\n'.join(block(lang) for lang in LANGUAGES)
    return (
        'You write a plain-language comprehension checkpoint for one clause of a '
        'student bursary agreement. Return STRICT JSON only, no prose, shaped as '
        '{"en": CHECKPOINT, "ms": CHECKPOINT, "ta": CHECKPOINT} where each '
        'CHECKPOINT is {"tag": str, "plain": str, "question": str, '
        '"options": [str, str, str], "correct": 0|1|2, "why": str}. Exactly one '
        'option is correct (its index is "correct"); the same index must be '
        'correct in all three languages. Keep it warm and simple; never invent '
        'terms the clause does not contain.\n\nCLAUSE:\n' + clause_text
    )


def _parse_quiz_json(raw):
    text = (raw or '').strip()
    # Tolerate a ```json fence around the payload.
    if text.startswith('```'):
        text = re.sub(r'^```[a-zA-Z]*\n?', '', text)
        text = re.sub(r'\n?```$', '', text).strip()
    try:
        data = json.loads(text)
    except (ValueError, TypeError):
        raise ContractsError('quiz_bad_json')
    if not isinstance(data, dict):
        raise ContractsError('quiz_bad_json')
    return data


def _quiz_payload_valid(payload):
    """Structural check (the Q2 shape): 3 options, a correct index in 0..2."""
    if not isinstance(payload, dict):
        return False
    options = payload.get('options')
    if not isinstance(options, list) or len(options) != 3:
        return False
    if not all(isinstance(o, str) and o.strip() for o in options):
        return False
    return payload.get('correct') in (0, 1, 2)


def generate_quiz(clause, *, model=None):
    """Generate the clause's en/ms/ta quiz via Gemini, validate the structure,
    and save it to the (draft) clause. Billable → call on demand only. Only a top-level
    (level 0) clause may carry a quiz; it covers the clause's whole subtree."""
    _require_draft(clause.template)
    if clause.level != 0:
        raise ContractsError('quiz_not_top_level')
    model = model or getattr(settings, 'CONTRACT_QUIZ_MODEL', 'gemini-2.5-pro')
    raw = _gemini_generate(_build_quiz_prompt(clause), model)
    data = _parse_quiz_json(raw)
    payloads = {}
    for lang in LANGUAGES:
        payload = data.get(lang) or {}
        if payload and not _quiz_payload_valid(payload):
            raise ContractsError('quiz_invalid', lang)
        payloads[lang] = payload
    if not payloads.get('en'):
        raise ContractsError('quiz_invalid', 'en missing')
    clause.quiz_en = payloads['en']
    clause.quiz_ms = payloads['ms']
    clause.quiz_ta = payloads['ta']
    clause.quiz_generated_model = model
    clause.is_quiz_candidate = True
    clause.save(update_fields=[
        'quiz_en', 'quiz_ms', 'quiz_ta', 'quiz_generated_model', 'is_quiz_candidate',
    ])
    return clause


# ─────────────────────────────────────────────────────────────────────────────
# Word import (populate a draft's clauses from an author's .docx — Sprint 4)
#
# The uploaded file is a POPULATE CONVENIENCE, never retained as a legal artefact:
# segment_docx returns a proposed [{heading, body}] list the author reviews, and only
# the reviewed structured clauses (saved via replace_clauses) are the source of truth.
# ─────────────────────────────────────────────────────────────────────────────
def _extract_docx_text(data):
    """Plain text of a .docx (paragraphs joined by newlines). Raises on an
    unreadable or empty document (the FE degrades to hand-editing)."""
    try:
        import docx
    except ImportError:
        raise ContractsError('docx_unavailable')
    import io as _io
    try:
        document = docx.Document(_io.BytesIO(data))
    except Exception:
        raise ContractsError('docx_unreadable')
    paragraphs = [p.text.strip() for p in document.paragraphs]
    text = '\n'.join(p for p in paragraphs if p)
    if not text.strip():
        raise ContractsError('docx_empty')
    return text


def _para_numbering(para):
    """``(numId, ilvl)`` from a paragraph's direct list formatting, or ``(None, None)``.
    Word's visible numbers (``1.``, ``1.1``, ``i)``) come from this list definition, NOT
    from ``paragraph.text`` — so this is how the structure is actually read. A missing
    ``ilvl`` element means level 0."""
    try:
        numPr = para._p.pPr.numPr
    except AttributeError:
        return None, None
    if numPr is None:
        return None, None
    numId = numPr.numId.val if numPr.numId is not None else None
    ilvl = numPr.ilvl.val if numPr.ilvl is not None else None
    return numId, ilvl


def _docx_structure(data):
    """Deterministically parse a Word ``.docx`` into our clause hierarchy using the
    document's OWN heading styles + list levels — no AI guess. Returns
    ``{'title': str, 'preamble': str, 'clauses': [{heading, body, level}]}`` or ``None``
    when the document has no usable heading structure (an unstyled / hand-numbered doc),
    so the caller can fall back to Gemini segmentation.

    The old text-only path lost structure because Word GENERATES the ``1.`` / ``1.1`` /
    ``i)`` labels from the style + list definition — they are not in ``paragraph.text``.
    Here we read that definition directly:

    * ``Title`` style                     → the agreement title.
    * a plain paragraph before clause 1   → preamble / parties recital.
    * a ``Heading N`` paragraph           → a clause; its level is the paragraph's list
      level (``ilvl``) when present, else ``N-1`` (Heading 1→0, Heading 2→1, Heading 3→2).
    * a plain paragraph carrying a list    → an item nested one level UNDER the current
      heading (so a definition/obligation list renders ``i) ii) …``).
    * a plain un-numbered paragraph after a clause → continuation body of that clause.

    Levels pass through :func:`normalise_levels` (the no-skip guard)."""
    try:
        import docx
    except ImportError:
        raise ContractsError('docx_unavailable')
    import io as _io
    try:
        document = docx.Document(_io.BytesIO(data))
    except Exception:
        raise ContractsError('docx_unreadable')

    title = ''
    preamble = []
    clauses = []
    current_heading_level = 0
    saw_clause = False
    for para in document.paragraphs:
        text = (para.text or '').strip()
        if not text:
            continue
        style = (para.style.name if para.style else '') or ''
        numId, ilvl = _para_numbering(para)
        if style == 'Title':
            if not title:
                title = text
            continue
        if style.startswith('Heading'):
            if ilvl is not None:
                level = ilvl
            else:
                m = re.search(r'(\d+)', style)
                level = (int(m.group(1)) - 1) if m else 0
            level = max(0, min(MAX_CLAUSE_LEVEL, level))
            # A short Heading is a clause TITLE; a long one is a full sub-clause the author
            # styled as a Heading (common) — that text is clause BODY, not a 255-char title.
            if len(text) > HEADING_TITLE_MAX:
                clauses.append({'heading': '', 'body': text, 'level': level})
            else:
                clauses.append({'heading': text, 'body': '', 'level': level})
            current_heading_level = level
            saw_clause = True
        elif numId is not None:
            level = min(MAX_CLAUSE_LEVEL, current_heading_level + 1) if saw_clause else 0
            clauses.append({'heading': '', 'body': text, 'level': level})
            saw_clause = True
        elif not saw_clause:
            preamble.append(text)
        elif clauses:
            last = clauses[-1]
            last['body'] = (last['body'] + '\n\n' + text) if last['body'] else text

    # Only usable if the document actually carried headings; otherwise it is an unstyled
    # doc and the AI fallback will read it better than this structural walk would.
    if not clauses or not any(c['heading'] for c in clauses):
        return None
    levels = normalise_levels([c['level'] for c in clauses])
    for c, lv in zip(clauses, levels):
        c['level'] = lv
    return {'title': title, 'preamble': '\n\n'.join(preamble).strip(), 'clauses': clauses}


def _build_segment_prompt(text):
    return (
        'You segment the plain text of a bursary-agreement document into its numbered '
        'clauses. Return STRICT JSON only, no prose: a list '
        '[{"heading": str, "body": str, "level": 0|1|2}, ...] in document order. Each heading '
        "is the clause's short title; each body is its plain-text content (keep paragraphs as "
        'blank-line-separated text). "level" is the nesting depth read from the numbering: '
        '0 for a main clause (1., 2., 3.), 1 for a sub-clause (1.1, 2.3), 2 for a sub-sub-clause '
        '(i), (ii), (a), (b)). A sub-clause must follow its parent; never skip a level. '
        'Do NOT invent, summarise, or reword — copy the wording. Drop page headers/footers, '
        'signature blocks and the title.\n\nDOCUMENT:\n' + text
    )


def _parse_segments(raw):
    text = (raw or '').strip()
    if text.startswith('```'):
        text = re.sub(r'^```[a-zA-Z]*\n?', '', text)
        text = re.sub(r'\n?```$', '', text).strip()
    try:
        data = json.loads(text)
    except (ValueError, TypeError):
        raise ContractsError('segmentation_failed')
    if not isinstance(data, list) or not data:
        raise ContractsError('segmentation_failed')
    clauses = []
    for item in data:
        if not isinstance(item, dict):
            raise ContractsError('segmentation_failed')
        heading = (item.get('heading') or '').strip()
        body = (item.get('body') or '').strip()
        if not (heading or body):
            continue
        clauses.append({'heading': heading, 'body': body, 'level': item.get('level', 0)})
    if not clauses:
        raise ContractsError('segmentation_failed')
    # Enforce the no-skip rule on the proposed levels (the model can drift); the author still
    # reviews the result before it is saved.
    levels = normalise_levels([c['level'] for c in clauses])
    for c, lv in zip(clauses, levels):
        c['level'] = lv
    return clauses


def segment_docx(data, *, model=None):
    """Propose a draft's structure from an author's ``.docx``. Tries the DETERMINISTIC
    parse of the document's own heading/list numbering first (:func:`_docx_structure`) —
    which reads the real ``1.`` / ``1.1`` / ``i)`` hierarchy Word generates — and only an
    unstyled document falls back to Gemini segmentation of its flat text.

    Returns ``{'title': str, 'preamble': str, 'clauses': [{heading, body, level}]}`` — the
    PROPOSAL only; nothing is saved and the upload is not retained (the caller reviews it
    and, on confirm, calls replace_clauses / update_config). Hand-typed bracket placeholders
    are converted to ``{{tokens}}``. Raises ContractsError (docx_unreadable / docx_empty /
    segmentation_failed / …) so the FE can degrade to hand-editing."""
    structured = _docx_structure(data)
    if structured is None:
        text = _extract_docx_text(data)
        model = model or getattr(settings, 'CONTRACT_QUIZ_MODEL', 'gemini-2.5-pro')
        raw = _gemini_generate(_build_segment_prompt(text), model)
        structured = {'title': '', 'preamble': '', 'clauses': _parse_segments(raw)}
    structured['preamble'] = _brackets_to_vars(structured.get('preamble', ''))
    for c in structured['clauses']:
        c['heading'] = _brackets_to_vars(c['heading'])
        c['body'] = _brackets_to_vars(c['body'])
    return structured


# ─────────────────────────────────────────────────────────────────────────────
# Deploy validation (T / C / Q / S / P rules + W warnings)
# ─────────────────────────────────────────────────────────────────────────────
def _expected_row_total(pathway, variant):
    """Mirror award.py: what the award (and therefore the schedule total) must be
    for a pathway/variant. STPM fresh 3000, STPM continuing 1000, else 2000."""
    if pathway == 'stpm':
        return award._STPM_CONTINUING_AMOUNT if variant == 'continuing' else award._STPM_AMOUNT
    return award._DEFAULT_AMOUNT


def validate_for_deployment(template):
    """Return a ValidationResult of error + warning codes. Errors block
    deployment; warnings are advisory (surfaced on the deploy panel)."""
    errors = []
    warnings = []
    clauses = list(template.clauses.all().order_by('order'))
    rows = list(template.schedule_rows.all())

    # ── T: template-level ─────────────────────────────────────────────────────
    if not (template.version and template.counterparty_name
            and template.counterparty_title and template.counterparty_nric):
        errors.append('T1')
    if not (template.vetted_by_name and template.vetted_on
            and template.vetting_attested_by_email):
        errors.append('T2')

    # ── C: clauses ────────────────────────────────────────────────────────────
    orders = [c.order for c in clauses]
    if not clauses or orders != list(range(1, len(clauses) + 1)):
        errors.append('C1')
    english_complete = bool(
        template.title_en and template.preamble_en and template.progress_standard_en
    ) and all(c.heading_en and c.body_en for c in clauses)
    if not clauses or not english_complete:
        errors.append('C2')

    # ── Q: quiz ───────────────────────────────────────────────────────────────
    candidates = [c for c in clauses if c.is_quiz_candidate]
    if not candidates:
        errors.append('Q1')
    for c in candidates:
        if not _quiz_payload_valid(c.quiz_en):
            errors.append('Q2')
    for c in clauses:
        if not c.is_quiz_candidate and (c.quiz_en or c.quiz_ms or c.quiz_ta):
            errors.append('Q3')  # a question can't outlive its clause
    for c in candidates:
        base = c.quiz_en.get('correct') if isinstance(c.quiz_en, dict) else None
        for lang in ('ms', 'ta'):
            payload = getattr(c, f'quiz_{lang}')
            if payload:
                if not _quiz_payload_valid(payload) or payload.get('correct') != base:
                    errors.append('Q4')

    # ── S: schedule ───────────────────────────────────────────────────────────
    by_key = {(r.pathway, r.variant): r for r in rows}
    if ('default', '') not in by_key:
        errors.append('S1')
    seen = set()
    for r in rows:
        key = (r.pathway, r.variant)
        shape_ok = (
            r.monthly_amount and r.monthly_amount > 0
            and 1 <= r.start_month <= 12
            and isinstance(r.paid_offsets, list) and r.paid_offsets
            and all(isinstance(o, int) and o >= 0 for o in r.paid_offsets)
            and len(set(r.paid_offsets)) == len(r.paid_offsets)
        )
        if not shape_ok or key in seen:
            errors.append('S2')
        seen.add(key)
        if r.total not in award.ALLOWED_AMOUNTS:
            errors.append('S3')
        if r.total != _expected_row_total(r.pathway, r.variant):
            errors.append('S4')

    # ── P: v1 fence ───────────────────────────────────────────────────────────
    if template.parent_role == 'minor_only' or template.witness_policy == 'required':
        errors.append('P1')

    # ── W: warnings (advisory) ────────────────────────────────────────────────
    if template.parent_role == 'co_signer_all':
        blob = ' '.join([template.preamble_en] + [c.body_en for c in clauses]).lower()
        if 'guarantor' in blob or 'surety' in blob:
            warnings.append('W1')
    if 'ms' not in template.languages_available:
        warnings.append('W2')
    if 'ta' not in template.languages_available:
        warnings.append('W2')
    rm_literal = re.compile(r'RM\s?\d')
    if any(rm_literal.search(c.body_en or '') for c in clauses):
        warnings.append('W3')

    return ValidationResult(errors, warnings)


# ─────────────────────────────────────────────────────────────────────────────
# Lifecycle
# ─────────────────────────────────────────────────────────────────────────────
def submit_for_deployment(template, *, submitted_by_email=''):
    """draft → pending_deployment. Refuses unless validation passes."""
    if template.status != 'draft':
        raise ContractsError('not_draft')
    result = validate_for_deployment(template)
    if not result.ok:
        err = ContractsError('not_deployable')
        err.errors = result.errors
        raise err
    template.status = 'pending_deployment'
    template.submitted_by_email = submitted_by_email or ''
    template.submitted_by_at = timezone.now()
    template.save(update_fields=[
        'status', 'submitted_by_email', 'submitted_by_at', 'updated_at',
    ])
    return template


def revert_to_draft(template):
    """pending_deployment → draft (to edit further before deploying)."""
    if template.status != 'pending_deployment':
        raise ContractsError('not_pending')
    template.status = 'draft'
    template.submitted_by_email = ''
    template.submitted_by_at = None
    template.save(update_fields=[
        'status', 'submitted_by_email', 'submitted_by_at', 'updated_at',
    ])
    return template


@transaction.atomic
def deploy(template, *, deployed_by_email='', is_super=False):
    """pending_deployment → active (SUPER only). Atomically archives the org's
    previous active version — exactly one active template per org."""
    if not is_super:
        raise ContractsError('deploy_forbidden')
    if template.status != 'pending_deployment':
        raise ContractsError('not_pending')
    result = validate_for_deployment(template)
    if not result.ok:
        err = ContractsError('not_deployable')
        err.errors = result.errors
        raise err
    now = timezone.now()
    previous = (ContractTemplate.objects
                .select_for_update()
                .filter(organisation=template.organisation, status='active')
                .exclude(pk=template.pk))
    for old in previous:
        old.status = 'archived'
        old.archived_at = now
        old.save(update_fields=['status', 'archived_at', 'updated_at'])
    template.status = 'active'
    template.deployed_by_email = deployed_by_email or ''
    template.deployed_by_at = now
    template.save(update_fields=[
        'status', 'deployed_by_email', 'deployed_by_at', 'updated_at',
    ])
    return template


# ─────────────────────────────────────────────────────────────────────────────
# Readers (the seams bursary.py / payments.py cut over to in Sprint 2)
# ─────────────────────────────────────────────────────────────────────────────
def active_template_for(organisation):
    return (ContractTemplate.objects
            .filter(organisation=organisation, status='active')
            .order_by('-deployed_by_at', '-created_at')
            .first())


def template_for_application(application):
    """The template governing this application: the signed agreement's pinned
    template if any, else the owning org's active template."""
    agreement = getattr(application, 'bursary_agreement', None)
    if agreement is not None and agreement.template_id:
        return agreement.template
    org = getattr(application, 'owning_organisation', None)
    return active_template_for(org) if org is not None else None


def schedule_row_for(template, application):
    """The schedule row for this application's pathway: pathway + 'continuing'
    (via award._stpm_continuing), falling back to variant='' then ('default','')."""
    if template is None:
        return None
    pathway = (getattr(application, 'chosen_pathway', '') or '').strip().lower()
    variant = 'continuing' if award._stpm_continuing(application) else ''
    by_key = {(r.pathway, r.variant): r for r in template.schedule_rows.all()}
    for key in [(pathway, variant), (pathway, ''), ('default', '')]:
        if key in by_key:
            return by_key[key]
    return None


def is_paid_month(row, cohort_year, month, year=None):
    """Whether a calendar (year, month) is a paid month under ``row``. ``year``
    defaults to ``cohort_year`` (the common in-cohort-year call); the STPM
    schedule spans two calendar years, so a second-year month passes ``year``."""
    if row is None:
        return False
    if year is None:
        year = cohort_year
    offset = (year - cohort_year) * 12 + (month - row.start_month)
    return offset in set(row.paid_offsets or [])


_MONTHS = ('', 'January', 'February', 'March', 'April', 'May', 'June',
           'July', 'August', 'September', 'October', 'November', 'December')


def _row_label(row, locale):
    return getattr(row, f'label_{locale}', '') or row.label_en or row.pathway


def schedule_summary_text(row, locale='en'):
    """A one-line plain-text summary of a row (fills BursaryAgreement.payment_schedule)."""
    if row is None:
        return ''
    count = len(row.paid_offsets or [])
    start = _MONTHS[row.start_month] if 1 <= row.start_month <= 12 else ''
    amount = f'RM{int(row.monthly_amount)}' if row.monthly_amount == row.monthly_amount.to_integral() \
        else f'RM{row.monthly_amount}'
    start_phrase = f', starting {start}' if start else ''
    return f'{amount} per month × {count} payments{start_phrase} (total RM{int(row.total)}).'


def schedule_table(template, locale='en'):
    """Structured rows for rendering Schedule 1 (start, count, total, offsets)."""
    if template is None:
        return []
    table = []
    for row in template.schedule_rows.all():
        table.append({
            'pathway': row.pathway,
            'variant': row.variant,
            'label': _row_label(row, locale),
            'monthly_amount': row.monthly_amount,
            'start_month': row.start_month,
            'months': len(row.paid_offsets or []),
            'paid_offsets': list(row.paid_offsets or []),
            'total': row.total,
        })
    return table


def schedule_calendar(row, cohort_year, locale='en'):
    """Month-by-month rows for Schedule 1: every month from the first paid month
    through the last, each flagged paid (with amount) or a gap ('exam month — no
    payment'). Drives the rendered contract's payment table."""
    if row is None or not row.paid_offsets or not cohort_year:
        return []
    offsets = sorted(int(o) for o in row.paid_offsets)
    paid = set(offsets)
    calendar = []
    for off in range(0, offsets[-1] + 1):
        total = (row.start_month - 1) + off
        year = cohort_year + total // 12
        month = total % 12 + 1
        calendar.append({
            'label': f'{_MONTHS[month]} {year}',
            'paid': off in paid,
            'amount': row.monthly_amount if off in paid else None,
        })
    return calendar


def quiz_checkpoints(template, locale='en'):
    """The comprehension checkpoints for the candidate clauses, in order, with a
    per-clause English fallback when the requested locale isn't translated."""
    if template is None:
        return []
    checkpoints = []
    for clause in template.clauses.all().order_by('order'):
        if not clause.is_quiz_candidate:
            continue
        payload = getattr(clause, f'quiz_{locale}', None) or clause.quiz_en
        if payload:
            checkpoints.append(payload)
    return checkpoints


def resolve_locale(locale, template):
    """Clamp a requested locale to one the template actually offers (else 'en')."""
    if template is None:
        return 'en'
    return locale if locale in template.languages_available else 'en'


def _bold(escaped):
    """Canonical markdown-``**bold**`` → ``<b>…</b>`` transform, run on an ALREADY-ESCAPED
    string so the emphasis can never inject markup. Non-greedy, single-line: an unmatched
    ``**`` is left literal. The single source of truth for bold — ``bursary._bold`` (the
    signed-agreement render) delegates here so the preview and the PDF stay in step."""
    return re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', escaped or '')


def render_preview_html(template, locale='en'):
    """A minimal standalone HTML preview (PREVIEW banner + English-authoritative notice).
    Sample particulars only — never a real signed artefact. Author text is HTML-ESCAPED,
    and the preview renders the SAME hierarchical numbering (1. / 1.1 / i)) and ``**bold**``
    as the signed agreement; ``{{variable}}`` tokens are left visible (no student context in
    a preview). Served as inline HTML and, via ``?output=pdf``, as the preview PDF."""
    from django.utils.html import escape as _esc
    if template is None:
        return '<html><body><p>No template.</p></body></html>'
    loc = resolve_locale(locale, template)

    def rich(text):
        return _bold(_esc(text or ''))

    title = getattr(template, f'title_{loc}', '') or template.title_en
    preamble = getattr(template, f'preamble_{loc}', '') or template.preamble_en
    parts = [
        '<div style="background:#fde68a;padding:8px;text-align:center;">PREVIEW — not a signed agreement</div>',
        '<div style="background:#eff6ff;padding:6px;font-size:12px;">'
        'The English version is authoritative; other languages are courtesy translations.</div>',
        f'<h1>{rich(title)}</h1>',
        f'<p>{rich(preamble)}</p>',
    ]
    clauses = list(template.clauses.all().order_by('order'))
    numbers = clause_numbers([(getattr(c, 'level', 0) or 0) for c in clauses])
    for clause, number in zip(clauses, numbers):
        level = getattr(clause, 'level', 0) or 0
        heading = getattr(clause, f'heading_{loc}', '') or clause.heading_en
        body = getattr(clause, f'body_{loc}', '') or clause.body_en
        paras = ''.join(f'<p>{rich(p)}</p>' for p in (body or '').split('\n\n') if p.strip())
        head = f' <b>{rich(heading)}</b>' if heading else ''
        parts.append(
            f'<div style="margin-left:{level * 18}px;">'
            f'<h3 style="margin:8px 0 2px 0;">{_esc(number)}{head}</h3>{paras}</div>')
    footer = f'Version {_esc(template.version)}'
    if template.vetted_by_name and template.vetted_on:
        footer += f' — Vetted by {_esc(template.vetted_by_name)}, {_esc(str(template.vetted_on))}'
    parts.append(f'<hr><p style="font-size:12px;color:#666;">{footer}</p>')
    return '<html><body>' + ''.join(parts) + '</body></html>'
